import os
import re
import openai
import numpy as np
import pdfplumber
from docx import Document
from typing import List, Tuple, Dict  # <-- Import generic types from typing
from dotenv import load_dotenv
from config import OPENAI_API_KEY

################################################
# Configuration
################################################
openai.api_key = OPENAI_API_KEY
OPENAI_MODEL = "gpt-3.5-turbo"
PLACEHOLDER_PATTERN = r"\{\{(.*?)\}\}"
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
CHUNK_SIZE = 50  # Approximate chunk size in characters for splitting
USER_INSTRUCTIONS = (
    "Use a professional tone, summarize key points clearly, "
    "and maintain relevant context. Preserve the template's structure.\n\n"
    "If references are missing or insufficient, provide a best-effort summary "
    "or note that more data was not found."
)

################################################
# Agent Definitions
################################################


class DocumentIngestionAgent:
    """
    Agent responsible for loading raw documents from disk in various formats
    (TXT, MD, PDF) and returning raw text content.
    """

    def __init__(self, doc_dir: str):
        self.doc_dir = doc_dir

    def run(self) -> List[str]:
        """
        Scans the directory, reads each supported file, and returns a list of texts.
        """
        texts: List[str] = []
        for filename in os.listdir(self.doc_dir):
            path = os.path.join(self.doc_dir, filename)
            if filename.lower().endswith((".txt", ".md")):
                with open(path, "r", encoding="utf-8") as f:
                    texts.append(f.read())
            elif filename.lower().endswith(".pdf"):
                pdf_text = self._extract_text_from_pdf(path)
                if pdf_text.strip():
                    texts.append(pdf_text)
        return texts

    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a PDF file using pdfplumber.
        """
        text_content: List[str] = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        return "\n".join(text_content)


class ChunkingEmbeddingAgent:
    """
    Agent that chunks documents and creates embeddings for each chunk.
    This enables more granular retrieval rather than using the entire document as a single block.
    """

    def __init__(self, chunk_size: int = CHUNK_SIZE):
        self.chunk_size = chunk_size

    def run(self, docs: List[str]) -> Tuple[List[str], np.ndarray]:
        """
        Splits each doc into chunks of size `chunk_size`, creates embeddings,
        and returns (list_of_chunks, embeddings_array).
        """
        all_chunks: List[str] = []
        for doc_text in docs:
            start = 0
            while start < len(doc_text):
                end = start + self.chunk_size
                chunk = doc_text[start:end]
                all_chunks.append(chunk)
                start = end

        # Generate embeddings for each chunk
        embeddings = openai_embeddings(
            all_chunks
        )  # We'll define openai_embeddings below OR use the embedding_model
        return all_chunks, embeddings


class RetrievalAgent:
    """
    Agent that retrieves the most relevant chunks for a given query or placeholder
    from the vector store (embeddings).
    """

    def __init__(self, chunks: List[str], embeddings: np.ndarray):
        self.chunks = chunks
        self.embeddings = embeddings

    def run(self, query: str, top_k: int = 3) -> List[str]:
        """
        Returns the top_k most relevant chunks for the given query using cosine similarity.
        """
        # Use the same embedding model or method to encode the query
        query_embedding = openai_embeddings([query])  # or embedding_model.encode([...])
        query_embedding = query_embedding[0]  # just the single vector

        scores = np.dot(self.embeddings, query_embedding)
        top_indices = np.argsort(scores)[::-1][:top_k]
        return [self.chunks[i] for i in top_indices]


class LLMContentAgent:
    """
    Agent that, given a placeholder label, instructions, and relevant reference chunks,
    calls the LLM to generate the fill-in text.
    """

    def __init__(self, model_name=OPENAI_MODEL):
        self.model_name = model_name

    def run(
        self, placeholder: str, instructions: str, references: List[str], max_retries=2
    ) -> str:
        """
        Generate content for a given placeholder using the provided references and instructions.
        """
        system_prompt = (
            "You are a helpful assistant that uses the provided references "
            "to produce well-structured, contextually relevant text matching the instructions."
        )

        user_prompt = f"""
Placeholder: {placeholder}
Instructions: {instructions}

Relevant References:
{chr(10).join(references)}

Please produce a coherent piece of text that fits into the template section.
If references are insufficient, provide a best-effort summary or note that more data was not found.
"""

        for attempt in range(max_retries):
            try:
                response = openai.ChatCompletion.create(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=0.7,
                    max_tokens=500,
                )
                content = response.choices[0].message.content.strip()
                if content:
                    return content
            except Exception as e:
                print(f"Error calling LLM (attempt {attempt+1}): {e}")
        return f"[Warning: No suitable content generated for {placeholder}]"


class TemplateFillingAgent:
    """
    Agent that loads a DOCX template, finds all placeholders, and replaces them
    with the generated content while preserving formatting.
    """

    def __init__(self, template_path: str):
        self.template_path = template_path

    def load_template(self) -> Document:
        """
        Load the DOCX template.
        """
        return Document(self.template_path)

    def extract_placeholders(self, doc: Document) -> List[str]:
        """
        Extract all placeholders ({{name}}) from the docx paragraphs and tables.
        """
        placeholders = set()
        # Check paragraphs
        for p in doc.paragraphs:
            matches = re.findall(PLACEHOLDER_PATTERN, p.text)
            for m in matches:
                placeholders.add(m.strip())
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(PLACEHOLDER_PATTERN, cell.text)
                    for m in matches:
                        placeholders.add(m.strip())
        return list(placeholders)

    def replace_placeholder_in_paragraph(
        self, paragraph, placeholder: str, new_text: str
    ):
        """
        Replace a placeholder in a single paragraph.
        Advanced users may need to manipulate runs for complex formatting.
        """
        old_text = paragraph.text
        placeholder_pattern = "{{" + placeholder + "}}"
        if placeholder_pattern in old_text:
            paragraph.text = old_text.replace(placeholder_pattern, new_text)

    def replace_placeholders_in_doc(
        self, doc: Document, placeholder_to_text: Dict[str, str]
    ) -> Document:
        """
        Replace placeholders in paragraphs and table cells.
        """
        # Replace in paragraphs
        for p in doc.paragraphs:
            for ph, val in placeholder_to_text.items():
                self.replace_placeholder_in_paragraph(p, ph, val)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for ph, val in placeholder_to_text.items():
                        placeholder_pattern = "{{" + ph + "}}"
                        if placeholder_pattern in cell.text:
                            cell.text = cell.text.replace(placeholder_pattern, val)
        return doc


################################################
# Example Helper for Creating Embeddings
################################################


def openai_embeddings(texts: List[str]) -> np.ndarray:
    """
    Simple placeholder function for generating embeddings (using an LLM or
    another approach). You could replace it with your own logic, or directly
    use the sentence_transformers/embedding_model below.
    """
    # If you're using sentence-transformers globally, you can do:
    #   return embedding_model.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    #
    # If you're using OpenAI for embeddings:
    #   - call openai.Embedding.create(...) for each text
    #
    # For demo, we'll just return random vectors (not recommended in real usage).
    # This is to illustrate that you can use a typed function.
    rng = np.random.default_rng()
    embedding_dim = 384  # typical dimension for e.g. "all-MiniLM-L6-v2"
    return rng.random((len(texts), embedding_dim), dtype=np.float32)


################################################
# Orchestration (main pipeline)
################################################


def main():
    # 1. Document ingestion
    docs_dir = "input_docs"  # Directory with TXT, MD, PDF files
    ingestion_agent = DocumentIngestionAgent(docs_dir)
    raw_texts = ingestion_agent.run()

    if not raw_texts:
        print(
            "No input documents found. Please place your PDF/TXT/MD files in 'input_docs/'."
        )
        return

    # 2. Chunking and Embedding
    chunking_agent = ChunkingEmbeddingAgent(chunk_size=CHUNK_SIZE)
    chunks, embeddings = chunking_agent.run(raw_texts)

    # 3. Template filling setup
    template_path = "template_1.docx"  # The user-provided DOCX template
    filling_agent = TemplateFillingAgent(template_path)

    # Load the DOCX template
    try:
        doc = filling_agent.load_template()
    except Exception as e:
        print(f"Error loading template at '{template_path}': {e}")
        return

    # Extract placeholders
    placeholders = filling_agent.extract_placeholders(doc)
    if not placeholders:
        print(
            "No placeholders found in the template. Make sure they are in the form {{PlaceholderName}}."
        )
        return

    # 4. For each placeholder: retrieve context and generate text via LLM
    retrieval_agent = RetrievalAgent(chunks, embeddings)
    llm_agent = LLMContentAgent(model_name=OPENAI_MODEL)

    placeholder_to_text: Dict[str, str] = {}
    for ph in placeholders:
        # a) Retrieve top chunks relevant to the placeholder
        top_chunks = retrieval_agent.run(ph, top_k=3)

        # If top_chunks are empty or only whitespace, provide fallback
        if not top_chunks or all(not chunk.strip() for chunk in top_chunks):
            top_chunks = [
                "[No references found in the input documents for this topic. "
                "Please provide a generic or placeholder summary if possible.]"
            ]

        # b) Generate content with the LLM (which now has fallback instructions)
        content = llm_agent.run(ph, USER_INSTRUCTIONS, top_chunks)
        placeholder_to_text[ph] = content

    # 5. Fill the template
    filled_doc = filling_agent.replace_placeholders_in_doc(doc, placeholder_to_text)

    # 6. Save the final DOCX
    output_path = "final_output.docx"
    filled_doc.save(output_path)
    print(f"Document generated and saved to: {output_path}")


if __name__ == "__main__":
    main()
